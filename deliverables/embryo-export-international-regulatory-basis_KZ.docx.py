from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT="/home/user/Dr-Ha--agent/deliverables/embryo-export-international-regulatory-basis_KZ.docx"
FN="Times New Roman"
GREEN=RGBColor(0x1F,0x51,0x32); DARK=RGBColor(0x22,0x22,0x22); GRAY=RGBColor(0x5A,0x5A,0x5A); WHITE=RGBColor(0xFF,0xFF,0xFF)

doc=Document()
s=doc.sections[0]; s.top_margin=Cm(1.7); s.bottom_margin=Cm(1.4); s.left_margin=Cm(2.1); s.right_margin=Cm(2.1)
st=doc.styles['Normal']; st.font.name=FN; st.font.size=Pt(10); st.element.rPr.rFonts.set(qn('w:eastAsia'),FN)

def font(run,size=10,color=DARK,bold=False):
    run.font.size=Pt(size); run.font.color.rgb=color; run.font.bold=bold; run.font.name=FN
    rf=run._element.get_or_add_rPr().get_or_add_rFonts(); rf.set(qn('w:eastAsia'),FN); rf.set(qn('w:cs'),FN)
def P(after=4,before=0,line=1.18,align=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_paragraph(); f=p.paragraph_format; f.space_after=Pt(after); f.space_before=Pt(before); f.line_spacing=line; f.alignment=align; return p
def A(p,t,**kw):
    r=p.add_run(t); font(r,**kw); return r
def rule(p,color="1F5132",sz=6):
    pPr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); bt=OxmlElement('w:bottom')
    bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),str(sz)); bt.set(qn('w:space'),'4'); bt.set(qn('w:color'),color); b.append(bt); pPr.append(b)
def shade_para(p,color):
    pPr=p._p.get_or_add_pPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),color); pPr.append(sh)
def shade_cell(c,color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),color); c._tc.get_or_add_tcPr().append(sh)
def cell_text(c,t,size=9,color=DARK,bold=False,align=WD_ALIGN_PARAGRAPH.LEFT):
    p=c.paragraphs[0]; p.alignment=align; p.paragraph_format.space_after=Pt(1); p.paragraph_format.space_before=Pt(1)
    font(p.add_run(t),size,color,bold)

# Title
p=P(2,align=WD_ALIGN_PARAGRAPH.CENTER); A(p,"Ірі қара мал эмбриондарын халықаралық саудалаудың нормативтік негізі",size=16,color=GREEN,bold=True)
p=P(8,align=WD_ALIGN_PARAGRAPH.CENTER); A(p,"Regulatory Basis for International Trade in Bovine Embryos",size=10.5,color=GRAY); rule(p)

p=P(10); A(p,"Бұл құжат Корея Республикасынан шығарылатын ірі қара мал (мысалы, голштин сүтті бағыттағы) эмбриондарын экспорттау кезінде импорттаушы елдің мемлекеттік органдарына (ауыл шаруашылығы министрлігі / ветеринариялық қызмет) ұсынуға арналған халықаралық ветеринариялық-санитариялық нормативтік негізді баяндайды. Ол нақты бір елге шектелмейді және барлық импорттаушы елдерге бірдей қолданылады.",size=10)

# A
p=P(3); A(p,"A. Негізгі халықаралық келісім (құқықтық негіз)",size=12.5,color=GREEN,bold=True)
p=P(2); A(p,"Дүниежүзілік сауда ұйымының (WTO) «Санитариялық және фитосанитариялық шараларды қолдану жөніндегі келісімі» (SPS келісімі), 3-бабы және А қосымшасы (3-тармақ) — жануарлар денсаулығы саласындағы халықаралық стандарттарды белгілейтін орган ретінде ",size=10); A(p,"WOAH (бұрынғы OIE)",size=10,bold=True); A(p," ұйымын ресми түрде таниды.",size=10)

# B
p=P(4); A(p,"B. Қолданылатын техникалық стандарт — WOAH «Terrestrial Animal Health Code»",size=12.5,color=GREEN,bold=True)
p=P(4); A(p,"WOAH «Жер үсті жануарларының денсаулығы кодексі». Эмбриондарға тікелей қатысты баптар:",size=10)

tbl=doc.add_table(rows=1,cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.LEFT; tbl.autofit=False
for c,w in zip(tbl.rows[0].cells,[Cm(2.3),Cm(14.4)]): c.width=w
h=tbl.rows[0].cells; shade_cell(h[0],"1F5132"); shade_cell(h[1],"1F5132")
cell_text(h[0],"Бап",9,WHITE,True,WD_ALIGN_PARAGRAPH.CENTER); cell_text(h[1],"Атауы / Title",9,WHITE,True)
data=[("4.8","In vivo алынған эмбриондарды жинау және өңдеу (мал, жылқы)  ·  In vivo derived embryos"),
      ("4.9","In vitro өндірілген эмбриондарды жинау және өңдеу (мал)  ·  In vitro produced embryos"),
      ("4.10","Эмбриондарды микроманипуляциялау (жынысын анықтау)  ·  Micromanipulation of embryos")]
for i,(n,t) in enumerate(data):
    rc=tbl.add_row().cells
    if i%2==1: shade_cell(rc[0],"EEF3EF"); shade_cell(rc[1],"EEF3EF")
    for c,w in zip(rc,[Cm(2.3),Cm(14.4)]): c.width=w
    cell_text(rc[0],n,9.5,GREEN,True,WD_ALIGN_PARAGRAPH.CENTER); cell_text(rc[1],t,9,DARK)

p=P(2,before=6); A(p,"▪ Әр аурудың тарауындағы «ірі қара мал эмбриондарын импорттау жөніндегі ұсынымдар» баптары қоса қолданылады — мысалы: аусыл (FMD) 8.8-тарау, нодулярлық дерматит (LSD) 11.9-тарау, ірі қараның губкатәрізді энцефалопатиясы (BSE) 11.4-тарау, т.б.",size=10)
p=P(10); A(p,"▪ IETS (Халықаралық эмбриондық технологиялар қоғамы) нұсқаулығы: ",size=10); A(p,"ірі қара малдың in vivo эмбриондары — 1-санат (елеусіз тәуекел)",size=10,bold=True); A(p,". Эмбрион трансплантациясы ауру таратудың тәуекелі өте төмен әдіс ретінде халықаралық деңгейде танылған.",size=10)

# C
p=P(3); A(p,"C. Ескертпелер",size=12.5,color=GREEN,bold=True)
for t in [
 "WOAH — бұрынғы OIE ұйымы. Кодекстің бұрынғы «Appendix (қосымша)» нөмірлері қазіргі «Chapter (тарау)» нөмірлеріне сәйкес келеді (мысалы, бұрынғы Appendix 3.3.4 → қазіргі 4.8-тарау).",
 "CITES (жойылып бара жатқан түрлердің халықаралық саудасы туралы конвенция) қосымшалары жабайы әрі жойылып бара жатқан түрлерге арналған; үй жануары болып табылатын голштин сиыр эмбриондарына қолданылмайды.",
 "Нақты кедендік негіз — кодекстің өзі емес, экспорттаушы елдің ветеринариялық қызметі (Корея — APQA) мен импорттаушы ел арасында келісілген екіжақты «Ветеринариялық сертификат (Veterinary Health Certificate)»; жоғарыдағы WOAH баптары оның техникалық негізі.",
 "Тарау нөмірлері кодекстің басылымына (edition) байланысты өзгеруі мүмкін; дәйексөз келтірмес бұрын woah.org сайтындағы қолданыстағы басылым мен ағылшынша атауын тексеру ұсынылады."]:
    p=P(3); A(p,"▪ ",size=10,color=GREEN,bold=True); A(p,t,size=10)

# Bottom line
p=P(2,before=8); shade_para(p,"EEF3EF"); A(p,"  Қорытынды",size=11,color=GREEN,bold=True)
p=P(10); shade_para(p,"EEF3EF"); A(p,"  Эмбриондар WOAH кодексіне (4.8 · 4.9 · 4.10-тараулар) сәйкес жиналып, өңделсе және импорттаушы ел талап ететін ауруларға қатысты импорттық шарттар орындалып, екіжақты ветеринариялық сертификат берілсе — халықаралық саудада кедергі болмайды. Яғни кодексті сақтау — сауданың техникалық алғышарты, ал түпкілікті шарт — импорттаушы елдің ветеринариялық талаптары.  ",size=10.5,bold=True)

# Footer
p=P(0,before=6); rule(p,color="BBBBBB",sz=4)
p=P(0,line=1.05); A(p,"Шығарушы / Issuer: Genetics Co., Ltd. (Республика Корея).  Дереккөздер: WOAH Terrestrial Animal Health Code (қолданыстағы басылым); WTO SPS Agreement (Art. 3, Annex A). Анықтамалық құжат — дәйексөз кезінде соңғы басылым түпнұсқасын тексеру қажет.",size=8,color=GRAY)

doc.save(OUT); print("saved",OUT)
