# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def base_style(doc, font_name):
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(10.5)
    # set East Asian font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)


def add_title(doc, text, size=15, color="1F3864"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_heading(doc, text, size=12, color="1F3864"):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_body(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string("7F7F7F")
    return p


def add_defs_table(doc, rows, font_name):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (a, b) in enumerate(rows):
        r = table.add_row().cells
        r[0].text = a
        r[1].text = b
        # bold left col + bg on header row
        for run in r[0].paragraphs[0].runs:
            run.bold = True
        if i == 0:
            set_cell_bg(r[0], "D9E1F2")
            set_cell_bg(r[1], "D9E1F2")
    # widths
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(12.5)
    return table


def add_sign_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for a, b in rows:
        r = table.add_row().cells
        r[0].text = a
        r[1].text = b
        for run in r[0].paragraphs[0].runs:
            run.bold = True
        set_cell_bg(r[0], "F2F2F2")
    for row in table.rows:
        row.cells[0].width = Cm(6.0)
        row.cells[1].width = Cm(11.0)
    return table


def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BFBFBF')
    pbdr.append(bottom)
    pPr.append(pbdr)


# ============================================================
# KOREAN DOCUMENT
# ============================================================
def build_korean(path):
    doc = Document()
    base_style(doc, "Malgun Gothic")
    for s in doc.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

    add_title(doc, "KOREA HOLSTEIN IVF EMBRYO IMPORTING HEALTH REQUIREMENT", 13.5)
    add_title(doc, "대한민국 홀스타인 IVF(체외생산) 수정란 수입 위생요건서", 13)
    add_note(doc, "문서 형태: 협의용 초안 (Draft for bilateral negotiation) — 양국 수의당국(대한민국 농림축산검역본부 APQA ↔ 싱가포르 식품청 SFA) 협의·합의로 최종 확정. 법률·수의 검토 필수. WOAH(구 OIE) 육상동물위생규약 및 IETS Manual 기준 준거.")
    add_note(doc, "대상(Scope): 대한민국산 홀스타인(Bos taurus, Holstein) 체외생산수정란(IVF / In Vitro Produced, IVP). 한우(Hanwoo) 유전자원 비대상.")
    divider(doc)

    add_heading(doc, "정의 (Definitions)")
    add_defs_table(doc, [
        ("용어", "정의"),
        ("IVF 수정란 (IVF/IVP Embryo)", "홀스타인 소(Bos taurus)에서 채취한 난자를 체외에서 수정·배양하여 생산하고 동결·보존한 초기 발생단계의 배(胚)."),
        ("공란우 (Donor)", "난자(난모세포)를 제공한 홀스타인 암소."),
        ("수출국 / 수입국", "대한민국 (Republic of Korea) / 싱가포르 (Republic of Singapore)."),
        ("정부수의관 (Official Veterinarian)", "수출국 정부(농림축산검역본부)가 본 증명 발급 권한을 부여한 수의사."),
        ("승인 처리시설 (Approved facility)", "대한민국 정부기관으로부터 「가축인공수정 관련 정액·수정란 처리업 허가」를 받아 수정란을 생산·처리·보관하는 시설."),
        ("IETS Manual", "국제수정란이식학회(International Embryo Technology Society) 발간 매뉴얼(최신판)."),
        ("WOAH (구 OIE)", "세계동물보건기구, 육상동물위생규약(Terrestrial Animal Health Code) 발간 기관."),
    ], "Malgun Gothic")

    add_heading(doc, "제1조 (원산지국의 동물질병 상태)")
    add_body(doc, "1.1 본 IVF 수정란은 대한민국에서 채란·생산·처리·동결·보관되었다.")
    add_body(doc, "1.2 채란일 기준 대한민국은 WOAH 기준에 따라 우역(Rinderpest) 청정국이며, 소해면상뇌증(BSE)에 대해 WOAH 인정 위험도 지위를 보유한다.")
    add_body(doc, "1.3 WOAH 육상동물위생규약은 온전한 투명대를 가진 적절히 세척된 수정란을 통한 전염성 질병의 전파 위험을 무시할 수 있는 것으로 규정한다. 따라서 본 위생요건은 수출국의 개별 질병 발생상황에 의존하지 아니하며, 제4조의 IETS·WOAH 세척·검사 기준의 충족을 통하여 위생 안전성이 보장된다.", bold=True)
    add_note(doc, "(본 조항은 수정란 매개 질병전파에 관한 WOAH/IETS 국제 표준원칙에 근거하며, 세부사항은 양국 수의당국 합의로 확정.)")

    add_heading(doc, "제2조 (공란우 건강 요건)")
    add_body(doc, "2.1 공란우는 채란일에 임상적으로 건강하였고 전염성 질병의 임상증상을 보이지 않았다.")
    add_body(doc, "2.2 공란우는 승인 처리시설의 표준운영절차(SOP) 및 IETS Manual에 따라 관리되는 사육·채란 환경에서 유래하였다.")
    add_body(doc, "2.3 공란우는 소결핵병·소브루셀라병·소백혈병(EBL)·IBR/IPV 등에 대해 검역본부 인정 방법으로 음성 확인되었거나, 해당 시설·농장의 위생관리 조건으로 갈음된다.")
    add_note(doc, "(검사 항목·시기는 양국 합의로 확정. 수정란 매개 질병전파 위험은 제4조 세척원칙으로 1차 통제됨.)")
    add_body(doc, "2.4 채란 전후 관찰·건강기록은 처리시설에 보관되며 요청 시 수입국에 제공된다.")

    add_heading(doc, "제3조 (정액의 출처)")
    add_body(doc, "3.1 IVF 생산에 사용된 정액은 WOAH 기준 및 대한민국 정부 승인 인공수정센터(AI Centre) 또는 승인 정액처리시설에서 생산·관리되었다.")
    add_body(doc, "3.2 사용 정액은 제1조·제2조에 준하는 위생·질병 요건을 충족하는 공우로부터 채취되었다.")

    add_heading(doc, "제4조 (IVF 수정란의 생산·처리·세척 — IETS 핵심 원칙) ★")
    add_body(doc, "4.1 모든 IVF 수정란은 IETS Manual 및 WOAH 육상동물위생규약(체외생산수정란 관련 챕터)에 따라 채란·취급·수정·배양·처리·세척·동결되었다.")
    add_body(doc, "4.2 세척·검사 핵심 원칙 — 각 수정란은 다음을 모두 충족한다.", bold=True)
    add_body(doc, "   (a) 투명대(zona pellucida)가 온전하며 손상이 없다.")
    add_body(doc, "   (b) 50배 이상 입체현미경(stereomicroscope) 하에서 투명대 표면에 이물질(adherent material) 부착이 없다.")
    add_body(doc, "   (c) IETS Manual에 따라 1군 10개 이하, 최소 10회 연속 희석 세척(매회 100배 이상 희석, 매 세척 시 멸균 피펫 교체)을 실시하였고, 특정 질병 우려 시 트립신(trypsin) 처리 등 추가 처리를 실시하였다.")
    add_body(doc, "4.3 온전한 투명대를 가진 수정란을 IETS Manual 절차에 따라 세척한 경우, 수정란을 매개로 한 전염성 질병의 전파 위험은 무시할 수 있다(negligible)는 IETS·WOAH 과학적 원칙이 적용된다. 따라서 본 IVF 수정란은 국제 위생기준에 부합한다.")
    add_body(doc, "4.4 배양에 사용된 배지·혈청·생물학적 원료는 질병 전파 위험이 없도록 관리되었으며 관련 기록은 처리시설에 보관된다.")

    add_heading(doc, "제5조 (국가 허가 처리시설 생산 수정란의 인정) ★")
    add_body(doc, "5.1 본 IVF 수정란은 대한민국의 「축산법」에 따라 관할 행정관청으로부터 「가축인공수정 관련 정액·수정란 처리업(정액등처리업) 허가」를 받은 시설에서, 시설 책임수의사의 감독 하에 생산·처리·동결·보관되었다. 본 허가는 대한민국 정부의 허가체계 하에서 부여된 국가 허가에 해당한다.")
    add_body(doc, "5.2 이러한 국가 허가 보유 시설은 제4조의 IETS·WOAH 위생기준에 부합하는 채취·세척·처리·보관 체계를 갖춘 것으로 인정되므로, 싱가포르는 해당 허가를 보유한 시설에서 본 위생요건을 충족하여 생산된 홀스타인 IVF 수정란의 수입을 허용한다.")
    add_body(doc, "5.3 본 수정란을 생산한 허가 처리시설의 정보는 다음과 같다.")
    add_sign_table(doc, [
        ("허가 처리시설명 (Licensed facility)", ""),
        ("허가번호 (Licence No.)", ""),
        ("허가 종류 (Type of licence)", "가축인공수정 관련 정액·수정란 처리업 (정액등처리업) / Processing of semen and embryos for artificial insemination"),
        ("허가관청 (Competent authority)", "경기도 / 안성 (Gyeonggi Province / Anseong, Republic of Korea)"),
        ("허가일자 (Date of licence)", ""),
        ("소재지 (Address)", ""),
    ])
    add_note(doc, "(허가번호·허가일자·정확한 관청명은 허가증 원본을 기준으로 기재. 시설 등록·인정 절차는 양국 합의로 확정.)")

    add_heading(doc, "제6조 (보관·식별·봉인)")
    add_body(doc, "6.1 수정란은 멸균 스트로(straw)에 충전되어 액체질소(또는 그 증기) 중에 동결·보관되었다.")
    add_body(doc, "6.2 각 스트로에 공란우 ID, 부(種牡牛) ID, 채란/생산일자, 처리시설 코드, 수정란 등급·단계가 식별 가능하게 표시된다.")
    add_body(doc, "6.3 신품 또는 효과적으로 소독된 멸균 용기를 사용하였으며, 액체질소·동결보호제는 질병 전파 위험이 없다.")
    add_body(doc, "6.4 수송 컨테이너(cryotank)는 충전 후 정부수의관 책임 하에 봉인되며, 봉인번호는 수정란 목록에 기재되고 수입국 검역관만 개봉한다.")

    add_heading(doc, "제7조 (수출 증명)")
    add_body(doc, "7.1 첨부서류: 상업송장, 포장명세서, 원산지증명서, 동물검역증명서, 수정란 목록(공란우/부/스트로 ID·봉인번호).")
    add_body(doc, "7.2 위생요건 충족은 정부수의관이 증명하며, 최종 수입 허용은 싱가포르 식품청(SFA) 검역 결정에 따른다.")

    add_heading(doc, "증명문 (Attestation)")
    add_body(doc, "본인은 대한민국 농림축산검역본부 권한 정부수의관으로서, 위 홀스타인 IVF 수정란이 본 위생요건서 제1조 내지 제7조의 모든 조건을 충족함을 증명한다.", italic=True)
    add_sign_table(doc, [
        ("정부수의관 성명 (Name)", ""),
        ("직책 / 소속 (Title / Authority)", "농림축산검역본부 / Animal and Plant Quarantine Agency (APQA)"),
        ("증명일자 (Date)", ""),
        ("서명 (Signature)", ""),
        ("관인 (Official Stamp)", ""),
    ])

    doc.save(path)
    print("Saved:", path)


# ============================================================
# ENGLISH DOCUMENT
# ============================================================
def build_english(path):
    doc = Document()
    base_style(doc, "Calibri")
    for s in doc.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

    add_title(doc, "KOREA HOLSTEIN IVF EMBRYO IMPORTING HEALTH REQUIREMENT", 14)
    add_title(doc, "Veterinary Health Requirements for the Importation of Korean Holstein IVF (In Vitro Produced) Bovine Embryos into the Republic of Singapore", 11)
    add_note(doc, "Status: Draft for bilateral agreement between the veterinary authorities of the Republic of Korea (Animal and Plant Quarantine Agency, APQA) and the Republic of Singapore (Singapore Food Agency, SFA). Subject to legal and veterinary review. Prepared in accordance with the WOAH (formerly OIE) Terrestrial Animal Health Code and the IETS Manual.")
    add_note(doc, "Scope: Korean Holstein (Bos taurus) in vitro produced (IVF/IVP) bovine embryos. Hanwoo genetics are excluded.")
    divider(doc)

    add_heading(doc, "Definitions")
    add_defs_table(doc, [
        ("Term", "Definition"),
        ("IVF/IVP Embryo", "An early-stage Holstein (Bos taurus) embryo produced in vitro from collected oocytes and cryopreserved."),
        ("Donor", "The Holstein female providing oocytes."),
        ("Exporting / Importing country", "Republic of Korea / Republic of Singapore."),
        ("Official Veterinarian", "A veterinarian authorised by APQA to issue this certification."),
        ("Approved facility", "A facility licensed by the Government of Korea (MAFRA/APQA) under the licence for processing semen and embryos for artificial insemination to produce, process and store embryos."),
        ("IETS Manual", "The current Manual of the International Embryo Technology Society."),
        ("WOAH", "The World Organisation for Animal Health (formerly OIE)."),
    ], "Calibri")

    add_heading(doc, "Article 1 — Animal health status of the country of origin")
    add_body(doc, "1.1 The IVF embryos were produced, processed, frozen and stored in the Republic of Korea.")
    add_body(doc, "1.2 As of the date of collection, the Republic of Korea is free from Rinderpest and holds a WOAH-recognised risk status for BSE (controlled or negligible risk).")
    add_body(doc, "1.3 The WOAH Terrestrial Animal Health Code provides that the risk of transmission of infectious disease through an embryo with an intact zona pellucida that has been properly washed is negligible. Accordingly, these health requirements do not depend on the individual disease occurrence status of the exporting country; sanitary safety is ensured through compliance with the IETS/WOAH washing and examination standards set out in Article 4.", bold=True)
    add_note(doc, "(This Article is based on the WOAH/IETS international standard principle on embryo-mediated disease transmission; details to be finalised by bilateral agreement.)")

    add_heading(doc, "Article 2 — Health requirements for donors")
    add_body(doc, "2.1 Donors were clinically healthy and showed no clinical signs of transmissible disease on the day of oocyte collection.")
    add_body(doc, "2.2 Donors originated from a husbandry and collection environment managed in accordance with the standard operating procedures (SOPs) of the approved facility and the IETS Manual.")
    add_body(doc, "2.3 Donors tested negative, by methods recognised by APQA, for Bovine Tuberculosis, Brucellosis, Enzootic Bovine Leukosis and IBR/IPV (and others as agreed), or this is satisfied by the sanitary management conditions of the facility/farm.")
    add_note(doc, "(Test panel and timing to be finalised by bilateral agreement. Embryo-mediated transmission risk is primarily controlled by the washing principle of Article 4.)")
    add_body(doc, "2.4 Records of donor observation and health are retained at the facility and made available to the importing authority on request.")

    add_heading(doc, "Article 3 — Source of semen")
    add_body(doc, "3.1 Semen used for IVF production originated from an AI Centre approved by the Government of Korea in accordance with WOAH standards.")
    add_body(doc, "3.2 The semen was collected from sires meeting health and disease requirements equivalent to Articles 1 and 2.")

    add_heading(doc, "Article 4 — Production, processing and washing of IVF embryos (IETS core principle) *")
    add_body(doc, "4.1 All IVF embryos were collected, fertilised, cultured, processed, washed and frozen in accordance with the IETS Manual and the relevant chapter of the WOAH Terrestrial Animal Health Code governing in vitro produced bovine embryos.")
    add_body(doc, "4.2 Each embryo meets all of the following:", bold=True)
    add_body(doc, "   (a) the zona pellucida is intact and undamaged;")
    add_body(doc, "   (b) on examination under a stereomicroscope at no less than 50x magnification, the zona pellucida shows no adherent material;")
    add_body(doc, "   (c) the embryo was washed in accordance with the IETS Manual — i.e. in groups of no more than 10 embryos, through at least ten successive washes (each a dilution of 100-fold or more) using a fresh sterile pipette at each transfer; where a specific (e.g. certain viral) disease is of concern, trypsin treatment or other measures per the IETS Manual were additionally applied.")
    add_body(doc, "4.3 Pursuant to the scientific principle established by the IETS and WOAH, an embryo with an intact zona pellucida that has been washed in accordance with the IETS Manual carries a negligible risk of transmitting infectious disease. The embryos therefore comply with international sanitary standards.")
    add_body(doc, "4.4 Culture media, sera and other biological inputs were managed so as to present no risk of disease transmission; records are retained at the facility.")

    add_heading(doc, "Article 5 — Recognition of embryos produced at a nationally licensed facility *")
    add_body(doc, "5.1 The IVF embryos were produced, processed, frozen and stored at a facility holding a licence for the processing of semen and embryos for artificial insemination, granted by the competent authority under the Livestock Industry Act of the Republic of Korea, under the supervision of the facility's responsible veterinarian. This licence constitutes a national licence granted under the licensing system of the Government of the Republic of Korea.")
    add_body(doc, "5.2 A facility holding such a national licence is recognised as operating a collection, washing, processing and storage system compliant with the IETS/WOAH sanitary standards of Article 4. Accordingly, Singapore permits the importation of Holstein IVF embryos produced in compliance with these requirements at such a licensed facility.")
    add_body(doc, "5.3 The particulars of the licensed processing facility that produced these embryos are as follows.")
    add_sign_table(doc, [
        ("Licensed facility", ""),
        ("Licence No.", ""),
        ("Type of licence", "Processing of semen and embryos for artificial insemination"),
        ("Competent authority", "Gyeonggi Province / Anseong, Republic of Korea"),
        ("Date of licence", ""),
        ("Address", ""),
    ])
    add_note(doc, "(Licence number, date and exact authority name to be entered per the original licence certificate; facility registration/recognition procedure to be finalised by bilateral agreement.)")

    add_heading(doc, "Article 6 — Storage, identification and sealing")
    add_body(doc, "6.1 Embryos were loaded into sterile straws and frozen and stored in liquid nitrogen (or its vapour).")
    add_body(doc, "6.2 Each straw is legibly marked with the donor ID, sire ID, collection/production date, facility code, and embryo grade/stage.")
    add_body(doc, "6.3 New or effectively sterilised containers were used; the liquid nitrogen and cryoprotectants present no risk of disease transmission.")
    add_body(doc, "6.4 The export cryotank was sealed under the responsibility of the Official Veterinarian after loading; the seal number is recorded in the accompanying embryo list and the seal is to be broken only by the importing quarantine officer.")

    add_heading(doc, "Article 7 — Export certification")
    add_body(doc, "7.1 The consignment is accompanied by: Commercial Invoice, Packing List, Certificate of Origin, Veterinary Health Certificate, and Embryo List (with donor/sire/straw IDs and seal number).")
    add_body(doc, "7.2 Compliance is certified by the Official Veterinarian of the exporting country; final import clearance is subject to the decision of the Singapore Food Agency (SFA).")

    add_heading(doc, "Attestation")
    add_body(doc, "I, the undersigned Official Veterinarian authorised by the Animal and Plant Quarantine Agency of the Republic of Korea, certify that the Holstein IVF embryos described herein meet all the conditions set out in Articles 1 to 7 of these Veterinary Health Requirements.", italic=True)
    add_sign_table(doc, [
        ("Name of Official Veterinarian", ""),
        ("Title / Authority", "Animal and Plant Quarantine Agency (APQA), Republic of Korea"),
        ("Date", ""),
        ("Signature", ""),
        ("Official Stamp", ""),
    ])

    doc.save(path)
    print("Saved:", path)


if __name__ == "__main__":
    build_korean("/home/user/Dr-Ha--agent/exports/Korea_Holstein_IVF_Embryo_Health_Requirement_KO.docx")
    build_english("/home/user/Dr-Ha--agent/exports/Korea_Holstein_IVF_Embryo_Health_Requirement_EN.docx")
