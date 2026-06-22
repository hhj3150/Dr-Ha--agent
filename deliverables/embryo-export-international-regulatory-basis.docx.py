from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT="/home/user/Dr-Ha--agent/deliverables/embryo-export-international-regulatory-basis.docx"
KOR="맑은 고딕"; ENG="Calibri"
GREEN=RGBColor(0x1F,0x51,0x32); DARK=RGBColor(0x22,0x22,0x22); GRAY=RGBColor(0x60,0x60,0x60); WHITE=RGBColor(0xFF,0xFF,0xFF)

doc=Document()
s=doc.sections[0]; s.top_margin=Cm(1.7); s.bottom_margin=Cm(1.4); s.left_margin=Cm(2.1); s.right_margin=Cm(2.1)
st=doc.styles['Normal']; st.font.name=ENG; st.font.size=Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'),KOR)

def font(run,size=10,color=DARK,bold=False):
    run.font.size=Pt(size); run.font.color.rgb=color; run.font.bold=bold; run.font.name=ENG
    rpr=run._element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts(); rf.set(qn('w:eastAsia'),KOR)
def P(after=4,before=0,line=1.18,align=WD_ALIGN_PARAGRAPH.LEFT):
    p=doc.add_paragraph(); f=p.paragraph_format; f.space_after=Pt(after); f.space_before=Pt(before); f.line_spacing=line; f.alignment=align; return p
def A(p,t,**kw):
    r=p.add_run(t); font(r,**kw); return r
def rule(p,color="1F5132",sz=6):
    pPr=p._p.get_or_add_pPr(); b=OxmlElement('w:pBdr'); bt=OxmlElement('w:bottom')
    bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),str(sz)); bt.set(qn('w:space'),'4'); bt.set(qn('w:color'),color); b.append(bt); pPr.append(b)
def shade_para(p,color):
    pPr=p._p.get_or_add_pPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),color); pPr.append(sh)
def shade_cell(c,color):
    tcPr=c._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),color); tcPr.append(sh)
def cell_text(c,t,size=9,color=DARK,bold=False,align=WD_ALIGN_PARAGRAPH.LEFT):
    p=c.paragraphs[0]; p.alignment=align; p.paragraph_format.space_after=Pt(1); p.paragraph_format.space_before=Pt(1)
    r=p.add_run(t); font(r,size,color,bold)

# ===== Title =====
p=P(2,align=WD_ALIGN_PARAGRAPH.CENTER); A(p,"소(가축) 수정란 수출 국제규정 근거",size=19,color=GREEN,bold=True)
p=P(8,align=WD_ALIGN_PARAGRAPH.CENTER); A(p,"Regulatory Basis for International Trade in Bovine Embryos",size=10.5,color=GRAY)
rule(p)

p=P(10); A(p,"본 자료는 한국산 소 수정란(홀스타인 젖소 등)의 해외 수출 시 수입국 정부(농업부·검역당국) 심사에 인용·제출하기 위한 국제 동물위생 규정 근거이다. 특정 국가에 한정되지 않으며 모든 수입국에 공통 적용된다.",size=10,color=DARK)

# ===== A =====
p=P(3); A(p,"A. 상위 국제협정 (법적 근거)",size=12.5,color=GREEN,bold=True)
p=P(2); A(p,"WTO 「위생 및 식물위생 조치의 적용에 관한 협정(SPS 협정)」 제3조 및 부속서 A 제3항 — 동물위생 분야 국제표준 설정기관으로 ",size=10); A(p,"WOAH(구 OIE)",size=10,bold=True); A(p,"를 공식 지정한다.",size=10)
p=P(10); A(p,"WTO SPS Agreement, Article 3 & Annex A (para 3): designates WOAH (formerly OIE) as the reference standard-setting body for animal-health measures affecting international trade.",size=8.5,color=GRAY)

# ===== B =====
p=P(4); A(p,"B. 적용 기술기준 — WOAH 「육상동물위생규약」",size=12.5,color=GREEN,bold=True)
p=P(4); A(p,"WOAH Terrestrial Animal Health Code — 수정란 직접 근거 조문은 다음과 같다.",size=10)

tbl=doc.add_table(rows=1,cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.LEFT; tbl.autofit=False
tbl.columns[0].width=Cm(2.3); tbl.columns[1].width=Cm(14.4)
for c,w in zip(tbl.rows[0].cells,[Cm(2.3),Cm(14.4)]): c.width=w
h=tbl.rows[0].cells
shade_cell(h[0],"1F5132"); shade_cell(h[1],"1F5132")
cell_text(h[0],"장 Ch.",9,WHITE,True,WD_ALIGN_PARAGRAPH.CENTER); cell_text(h[1],"제목 / Title",9,WHITE,True)
data=[("4.8","체내(생체)수정란 채취·가공  ·  In vivo derived embryos (livestock & equids)"),
      ("4.9","체외생산수정란(IVF) 채취·가공  ·  In vitro produced embryos (livestock)"),
      ("4.10","수정란 미세조작 — 성감별 관련  ·  Micromanipulation of embryos")]
for i,(n,t) in enumerate(data):
    rc=tbl.add_row().cells
    if i%2==1: shade_cell(rc[0],"EEF3EF"); shade_cell(rc[1],"EEF3EF")
    for c,w in zip(rc,[Cm(2.3),Cm(14.4)]): c.width=w
    cell_text(rc[0],n,9.5,GREEN,True,WD_ALIGN_PARAGRAPH.CENTER); cell_text(rc[1],t,9,DARK)

p=P(2,before=6); A(p,"▪ 질병별 장의 ‘소 수정란 수입 권고(Recommendations for importation of bovine embryos)’ 조문 병행 적용 — 예: 구제역 8.8장, 럼피스킨병 11.9장, 소해면상뇌증(BSE) 11.4장 등.",size=10)
p=P(10); A(p,"▪ IETS(국제수정란기술학회) Manual: ",size=10); A(p,"체내수정란(소) = Category 1(위험 무시 가능)",size=10,bold=True); A(p," — 수정란 이식은 질병 전파 위험이 매우 낮은 방법으로 국제적으로 공인된다.",size=10)

# ===== C =====
p=P(3); A(p,"C. 유의사항",size=12.5,color=GREEN,bold=True)
for t in [
 "WOAH는 구 OIE(국제수역사무국)이며, 구 규약의 ‘Appendix(부속서)’는 현행 ‘Chapter(장)’에 대응한다(예: 구 Appendix 3.3.4 → 현 제4.8장).",
 "CITES(멸종위기종 국제거래협약)의 Appendix는 야생·멸종위기종을 위한 것으로, 가축인 홀스타인 젖소 수정란에는 적용되지 않는다.",
 "실제 통관 근거는 규약 자체가 아니라 수출국 검역당국(한국 APQA) ↔ 수입국이 합의한 양자 「수출검역증명서(Veterinary Health Certificate)」이며, 위 WOAH 조문은 그 기술적 토대이다.",
 "장 번호는 규약 ‘판(edition)’에 따라 변동될 수 있으므로, 인용 전 woah.org 현행판과 영문 제목을 함께 확인할 것을 권장한다."]:
    p=P(3); A(p,"▪ ",size=10,color=GREEN,bold=True); A(p,t,size=10)

# ===== Bottom line (shaded) =====
p=P(2,before=8); shade_para(p,"EEF3EF"); A(p,"  결론",size=11,color=GREEN,bold=True)
p=P(10); shade_para(p,"EEF3EF"); A(p,"  수정란을 WOAH 규약(제4.8 · 4.9 · 4.10장)대로 채취·가공하고, 수입국이 요구하는 질병별 수입조건을 충족하여 양자 수출검역증명서를 발급받으면 국제 교역에 문제가 없다. 즉 규약 준수는 교역의 기술적 전제이며, 최종 조건은 수입국의 검역요건이다.  ",size=10.5,color=DARK,bold=True)

# ===== Footer (no AI attribution) =====
p=P(0,before=6); rule(p,color="BBBBBB",sz=4)
p=P(0,line=1.05); A(p,"출처: WOAH Terrestrial Animal Health Code (현행판) · WTO SPS Agreement (Art. 3, Annex A).  본 자료는 참고용이며 인용 시 최신판 원문 확인을 요한다.",size=8,color=GRAY)

doc.save(OUT); print("saved", OUT)
